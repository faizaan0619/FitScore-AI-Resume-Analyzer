"""
resume_parser.py
Extracts raw text and lightweight structural signals from an uploaded resume file.
Supports PDF, DOCX, and TXT.
"""
import re
import io
import pdfplumber
import docx


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route to the correct extractor based on file extension."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        return _extract_from_pdf(file_bytes)
    elif ext == "docx":
        return _extract_from_docx(file_bytes)
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Please upload a PDF, DOCX, or TXT file.")


def _extract_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_from_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs]
    # Also pull table cell text (many resumes use tables for layout)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


# --- Structural signal detection -------------------------------------------------

SECTION_PATTERNS = {
    "contact_info": r"(email|phone|linkedin|github|@)",
    "summary": r"\b(summary|objective|profile)\b",
    "experience": r"\b(experience|employment|work history)\b",
    "education": r"\b(education|academic|degree|university|college)\b",
    "skills": r"\b(skills|technical skills|technologies|tools)\b",
    "projects": r"\b(projects|portfolio)\b",
}

EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3,5}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}")
LINK_RE = re.compile(r"(linkedin\.com|github\.com)/\S+", re.IGNORECASE)
BULLET_RE = re.compile(r"^[\s]*[•\-\*▪◦‣]\s+", re.MULTILINE)
DATE_RANGE_RE = re.compile(
    r"(19|20)\d{2}\s*(-|–|to)\s*((19|20)\d{2}|present|current)", re.IGNORECASE
)


def detect_sections(text: str) -> dict:
    lowered = text.lower()
    found = {}
    for section, pattern in SECTION_PATTERNS.items():
        found[section] = bool(re.search(pattern, lowered))
    return found


def structural_signals(text: str) -> dict:
    """Signals used by the ATS checker: contact details, bullets, dates, length."""
    word_count = len(text.split())
    return {
        "has_email": bool(EMAIL_RE.search(text)),
        "has_phone": bool(PHONE_RE.search(text)),
        "has_linkedin_or_github": bool(LINK_RE.search(text)),
        "bullet_count": len(BULLET_RE.findall(text)),
        "date_range_count": len(DATE_RANGE_RE.findall(text)),
        "word_count": word_count,
        "sections": detect_sections(text),
    }
